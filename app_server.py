from email.parser import BytesParser
from email.policy import default
from http import HTTPStatus
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
import json
from pathlib import Path
from tempfile import TemporaryDirectory
from urllib.parse import unquote

from docx import Document
from docx.shared import Pt

from gerar_controle import DEFAULT_OUTPUT_NAME, fill_document


BASE_DIR = Path(__file__).resolve().parent
MAX_UPLOAD_SIZE = 50 * 1024 * 1024


class AppHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(BASE_DIR), **kwargs)

    def do_GET(self):
        super().do_GET()

    def do_POST(self):
        path = self.path.split("?", 1)[0]

        if path == "/controle/gerar_docx":
            self.generate_controle_docx()
            return

        self.send_error(HTTPStatus.NOT_FOUND)

    def generate_controle_docx(self):
        try:
            filename, data, sheet_index = self.parse_xlsx_upload()
            with TemporaryDirectory() as temp_dir:
                temp_dir = Path(temp_dir)
                input_path = temp_dir / filename
                output_path = temp_dir / DEFAULT_OUTPUT_NAME
                input_path.write_bytes(data)
                fill_document(input_path, output_path, sheet_index=sheet_index)
                apply_docx_font(output_path)
                result = output_path.read_bytes()

            self.send_response(HTTPStatus.OK)
            self.send_header(
                "Content-Type",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            self.send_header(
                "Content-Disposition",
                "attachment; filename*=UTF-8''controle_preenchido.docx",
            )
            self.send_header("Content-Length", str(len(result)))
            self.end_headers()
            self.wfile.write(result)
        except Exception as exc:
            self.send_text(str(exc), status=HTTPStatus.BAD_REQUEST)

    def parse_xlsx_upload(self):
        content_type = self.headers.get("Content-Type", "")
        content_length = int(self.headers.get("Content-Length", "0") or "0")

        if content_length <= 0:
            raise ValueError("Nenhum arquivo foi enviado.")
        if content_length > MAX_UPLOAD_SIZE:
            raise ValueError("Arquivo muito grande. Limite: 50 MB.")
        if "multipart/form-data" not in content_type:
            raise ValueError("Formulario invalido.")

        body = self.rfile.read(content_length)
        mime_body = (
            f"Content-Type: {content_type}\r\nMIME-Version: 1.0\r\n\r\n".encode("utf-8")
            + body
        )
        message = BytesParser(policy=default).parsebytes(mime_body)

        filename = "planilha.xlsx"
        data = None
        sheet_index = 0

        for part in message.iter_parts():
            field_name = part.get_param("name", header="content-disposition")
            if field_name == "planilha":
                filename = safe_xlsx_filename(part.get_filename())
                data = part.get_payload(decode=True) or b""
                continue
            if field_name == "planilha_numero":
                raw_value = (part.get_payload(decode=True) or b"0").decode("utf-8", errors="ignore")
                try:
                    sheet_index = int(raw_value.strip())
                except ValueError:
                    sheet_index = 0
                sheet_index = max(0, min(sheet_index, 7))

        if not data:
            raise ValueError("Selecione uma planilha .xlsx.")

        return filename, data, sheet_index

    def send_json(self, data, status=HTTPStatus.OK, headers=None):
        content = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(content)))
        for name, value in (headers or {}).items():
            self.send_header(name, value)
        self.end_headers()
        self.wfile.write(content)

    def send_text(self, text, status=HTTPStatus.OK):
        content = str(text).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "text/plain; charset=utf-8")
        self.send_header("Content-Length", str(len(content)))
        self.end_headers()
        self.wfile.write(content)


def safe_xlsx_filename(name):
    cleaned = Path(name or "planilha.xlsx").name
    return cleaned if cleaned.lower().endswith(".xlsx") else "planilha.xlsx"


def apply_docx_font(docx_path):
    document = Document(docx_path)

    for style in document.styles:
        if getattr(style, "type", None) is None:
            continue
        if hasattr(style, "font"):
            style.font.name = "Arial"
            style.font.size = Pt(12)

    for paragraph in iter_docx_paragraphs(document):
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(12)

    document.save(docx_path)


def iter_docx_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from iter_table_paragraphs(table)
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for table in section.header.tables:
            yield from iter_table_paragraphs(table)
        for paragraph in section.footer.paragraphs:
            yield paragraph
        for table in section.footer.tables:
            yield from iter_table_paragraphs(table)


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from iter_table_paragraphs(nested_table)


def main():
    server = ThreadingHTTPServer(("127.0.0.1", 8770), AppHandler)
    print("App principal aberto em http://127.0.0.1:8770/app_navegador/index.html")
    server.serve_forever()


if __name__ == "__main__":
    main()
