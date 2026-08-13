import argparse
import html
import re
import shutil
import sys
from copy import deepcopy
from datetime import datetime, timedelta
from pathlib import Path
from tempfile import TemporaryDirectory
from zipfile import ZIP_DEFLATED, ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from openpyxl import load_workbook


BASE_DIR = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
TEMPLATE_DOCX = BASE_DIR / "modelo_controle.docx"
TEMPLATE_HTML = BASE_DIR / "modelo_controle.html"
DEFAULT_OUTPUT_NAME = "controle_preenchido.docx"
DEFAULT_HTML_OUTPUT_NAME = "controle_preenchido.html"

TYPE_LABELS = ["RESTR.", "LTS", "CONVAL", "LIC. COMP."]
PLACE_LABELS = ["CASA", "QUARTEL"]


def text(value):
    return "" if value is None else str(value).strip()


def html_text(value):
    escaped = html.escape(text(value), quote=True)
    return escaped.encode("ascii", "xmlcharrefreplace").decode("ascii")


def uppercase(value):
    return text(value).upper()


def to_datetime(value):
    if isinstance(value, datetime):
        return value
    if isinstance(value, str) and value.strip():
        for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d", "%d/%m/%Y"):
            try:
                return datetime.strptime(value.strip(), fmt)
            except ValueError:
                pass
    return None


def format_date(value):
    date = to_datetime(value)
    return date.strftime("%d/%m/%Y") if date else text(value)


def binary_options(value):
    selected = uppercase(value)
    if selected == "SIM":
        return "(X) SIM\n( ) NÃO"
    return "( ) SIM\n(X) NÃO"


def restriction_type(value):
    selected = uppercase(value).replace("REST", "RESTR").rstrip(".")
    aliases = {
        "RESTR": "RESTR.",
        "LTS": "LTS",
        "CONVAL": "CONVAL",
        "LIC. COMP": "LIC. COMP.",
        "LIC COMP": "LIC. COMP.",
    }
    selected = aliases.get(selected, aliases.get(selected.replace(".", ""), selected))
    return "\n".join(
        ("(X) " if label == selected else "( ) ") + label for label in TYPE_LABELS
    )


def place_options(value):
    selected = "CASA" if uppercase(value) == "CASA" else "QUARTEL"
    return "\n".join(
        ("(X) " if label == selected else "( ) ") + label for label in PLACE_LABELS
    )


def days_text(value):
    try:
        days = int(value)
    except (TypeError, ValueError):
        return text(value)
    return f"{days} DIA" if days == 1 else f"{days} DIAS"


def set_cell_text(cell, value):
    value = "" if value is None else str(value)
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.alignment = 1

    if paragraph.runs:
        first_run = paragraph.runs[0]
        for run in paragraph.runs:
            run.text = ""
    else:
        first_run = paragraph.add_run()

    lines = value.split("\n") or [""]
    first_run.text = lines[0]
    for line in lines[1:]:
        first_run.add_break()
        first_run.add_text(line)

    for extra_paragraph in cell.paragraphs[1:]:
        for run in extra_paragraph.runs:
            run.text = ""


def sheet_offset(sheet_index):
    try:
        index = int(sheet_index)
    except (TypeError, ValueError):
        index = 0
    return max(0, min(index, 7)) * 40


def cell_value(sheet, column, base_row, offset):
    return sheet[f"{column}{base_row + offset}"].value


def load_records(sheet, offset=0):
    records = []
    for row_number in range(10 + offset, 21 + offset):
        if not text(sheet[f"D{row_number}"].value):
            continue

        start = to_datetime(sheet[f"K{row_number}"].value)
        end = to_datetime(sheet[f"L{row_number}"].value)
        if end is None and start is not None and sheet[f"J{row_number}"].value not in (None, ""):
            end = start + timedelta(days=int(sheet[f"J{row_number}"].value) - 1)

        records.append(
            [
                uppercase(sheet[f"D{row_number}"].value),
                text(sheet[f"E{row_number}"].value),
                restriction_type(sheet[f"F{row_number}"].value),
                place_options(sheet[f"G{row_number}"].value),
                uppercase(sheet[f"H{row_number}"].value),
                uppercase(sheet[f"I{row_number}"].value),
                days_text(sheet[f"J{row_number}"].value),
                f"{format_date(start)} ATÉ {format_date(end)}",
                binary_options(sheet[f"M{row_number}"].value),
                binary_options(sheet[f"N{row_number}"].value),
            ]
        )
    return records


def fill_table(document, records):
    table = document.tables[0]
    template_row = deepcopy(table.rows[1]._tr)

    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)

    for record in records:
        table._tbl.append(deepcopy(template_row))
        new_row = table.rows[-1]
        for cell, value in zip(new_row.cells, record):
            set_cell_text(cell, value)


def normalized(value):
    value = value.replace("“", '"').replace("”", '"')
    return re.sub(r"\s+", " ", value).strip()


def set_paragraph_xml_text(text_nodes, value):
    if not text_nodes:
        return
    text_nodes[0].text = value
    for node in text_nodes[1:]:
        node.text = ""


def replace_split_paragraphs(docx_path, replacements):
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    ET.register_namespace("w", ns["w"])
    normalized_replacements = {
        normalized(old): replacement for old, replacement in replacements.items()
    }

    with TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)
        with ZipFile(docx_path) as source:
            source.extractall(temp_dir)

        changed = 0
        for xml_path in (temp_dir / "word").glob("*.xml"):
            tree = ET.parse(xml_path)
            root = tree.getroot()
            file_changed = False

            for paragraph in root.findall(".//w:p", ns):
                text_nodes = paragraph.findall(".//w:t", ns)
                if not text_nodes:
                    continue

                current = "".join(node.text or "" for node in text_nodes)
                replacement = normalized_replacements.get(normalized(current))
                if replacement is None or current == replacement:
                    continue

                set_paragraph_xml_text(text_nodes, replacement)
                changed += 1
                file_changed = True

            if file_changed:
                tree.write(xml_path, encoding="utf-8", xml_declaration=True)

        with ZipFile(docx_path, "w", ZIP_DEFLATED) as target:
            for path in temp_dir.rglob("*"):
                if path.is_file():
                    target.write(path, path.relative_to(temp_dir).as_posix())

    return changed


def fill_document(xlsx_path, output_path, sheet_index=0):
    if not TEMPLATE_DOCX.exists():
        raise FileNotFoundError(f"Modelo DOCX não encontrado: {TEMPLATE_DOCX}")

    workbook = load_workbook(xlsx_path, data_only=True)
    sheet = workbook["controle"] if "controle" in workbook.sheetnames else workbook.active
    offset = sheet_offset(sheet_index)
    records = load_records(sheet, offset)

    if not records:
        raise ValueError(f"Nenhum nome foi encontrado no intervalo D{10 + offset}:D{20 + offset}.")

    shutil.copyfile(TEMPLATE_DOCX, output_path)
    document = Document(output_path)
    fill_table(document, records)
    document.save(output_path)

    replacements = {
        'Pelotão: 3º CFO "A"': f"Pelotão: {text(sheet['E4'].value)}",
        "Período: 01 a 31 de maio de 2026": f"Período: {text(sheet['G4'].value)}",
        "HIGOR MACHADO MARQUES": text(cell_value(sheet, "G", 5, offset)),
        "ANDERSON ALVES SILVA": text(cell_value(sheet, "G", 6, offset)),
        "GEAZI DOS SANTOS RODRIGUES": text(cell_value(sheet, "G", 7, offset)),
        'Cad PM - Resp. Restr./Conval./LTS do 3º CFO "A"': text(sheet["F5"].value),
        "1º Ten PM - Oficial Responsável": text(sheet["F6"].value),
        "1º Tem PM - Cmt da 1ª Cia Es": text(sheet["F7"].value),
    }
    replacement_keys = list(replacements)
    replacements[replacement_keys[0]] = f"{replacement_keys[0].split(':', 1)[0]}: {text(cell_value(sheet, 'E', 4, offset))}"
    replacements[replacement_keys[1]] = f"{replacement_keys[1].split(':', 1)[0]}: {text(cell_value(sheet, 'G', 4, offset))}"
    replacements[replacement_keys[5]] = text(cell_value(sheet, "F", 5, offset))
    replacements[replacement_keys[6]] = text(cell_value(sheet, "F", 6, offset))
    replacements[replacement_keys[7]] = text(cell_value(sheet, "F", 7, offset))
    changed = replace_split_paragraphs(output_path, replacements)
    return len(records), changed


def read_html_template():
    if not TEMPLATE_HTML.exists():
        raise FileNotFoundError(f"Modelo HTML não encontrado: {TEMPLATE_HTML}")

    data = TEMPLATE_HTML.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def html_binary_options(value):
    selected = uppercase(value)
    yes = "(X)" if selected == "SIM" else "( )"
    no = "( )" if selected == "SIM" else "(X)"
    return f"{yes} SIM<br>{no} N&#195;O"


def html_restriction_type(value):
    selected = uppercase(value).replace("REST", "RESTR").rstrip(".")
    aliases = {
        "RESTR": "RESTR.",
        "LTS": "LTS",
        "CONVAL": "CONVAL",
        "LIC. COMP": "LIC. COMP.",
        "LIC COMP": "LIC. COMP.",
    }
    selected = aliases.get(selected, aliases.get(selected.replace(".", ""), selected))
    return "<br>".join(
        ("(X) " if label == selected else "( ) ") + html_text(label)
        for label in TYPE_LABELS
    )


def html_place_options(value):
    selected = "CASA" if uppercase(value) == "CASA" else "QUARTEL"
    return "<br>".join(
        ("(X) " if label == selected else "( ) ") + html_text(label)
        for label in PLACE_LABELS
    )


def html_period(start, end):
    return f"{html_text(format_date(start))} AT&#201; {html_text(format_date(end))}"


def load_html_records(sheet, offset=0):
    records = []
    for row_number in range(10 + offset, 21 + offset):
        if not text(sheet[f"D{row_number}"].value):
            continue

        start = to_datetime(sheet[f"K{row_number}"].value)
        end = to_datetime(sheet[f"L{row_number}"].value)
        if end is None and start is not None and sheet[f"J{row_number}"].value not in (None, ""):
            end = start + timedelta(days=int(sheet[f"J{row_number}"].value) - 1)

        records.append(
            [
                html_text(uppercase(sheet[f"D{row_number}"].value)),
                html_text(sheet[f"E{row_number}"].value),
                html_restriction_type(sheet[f"F{row_number}"].value),
                html_place_options(sheet[f"G{row_number}"].value),
                html_text(uppercase(sheet[f"H{row_number}"].value)),
                html_text(uppercase(sheet[f"I{row_number}"].value)),
                html_text(days_text(sheet[f"J{row_number}"].value)),
                html_period(start, end),
                html_binary_options(sheet[f"M{row_number}"].value),
                html_binary_options(sheet[f"N{row_number}"].value),
            ]
        )
    return records


def generated_cell(value):
    return (
        '<p style="margin:0;color:#000;font-family:&quot;Times New Roman&quot;,serif;'
        'font-size:12pt;line-height:1.15;text-align:center">'
        f"{value or '&nbsp;'}</p>"
    )


def replace_td_content(row_html, values):
    def repl(match):
        index = repl.index
        repl.index += 1
        if index >= len(values):
            return match.group(0)
        return f"{match.group(1)}{generated_cell(values[index])}{match.group(3)}"

    repl.index = 0
    return re.sub(r"(<td\b[^>]*>)([\s\S]*?)(</td>)", repl, row_html)


def replace_html_table(template, records):
    table_match = re.search(
        r"(<table\b[^>]*>[\s\S]*?<tr>[\s\S]*?</tr>)([\s\S]*?)(</table>)",
        template,
        flags=re.IGNORECASE,
    )
    if not table_match:
        raise ValueError("Tabela principal não encontrada no modelo HTML.")

    first_row_match = re.search(r"<tr>[\s\S]*?</tr>", table_match.group(2), flags=re.IGNORECASE)
    if not first_row_match:
        raise ValueError("Linha modelo da tabela não encontrada no modelo HTML.")

    template_row = first_row_match.group(0)
    rows = "\n".join(replace_td_content(template_row, record) for record in records)
    return (
        template[: table_match.start()]
        + table_match.group(1)
        + "\n"
        + rows
        + "\n"
        + table_match.group(3)
        + template[table_match.end() :]
    )


def replace_html_text(template, replacements):
    for old, new in replacements.items():
        if not new:
            continue
        template = template.replace(old, html_text(new))
        template = template.replace(html.escape(old), html_text(new))
    return template


def clean_control_html_template(template):
    header_patterns = [
        r"\s*<p\b[^>]*>[\s\S]*?N(?:&ordm;|º|\\u00ba)\s+do\s+Processo:[\s\S]*?</p>",
        r"\s*<p\b[^>]*>[\s\S]*?Interessado:[\s\S]*?</p>",
        r"\s*<p\b[^>]*>[\s\S]*?Assunto:[\s\S]*?Controle\s+de\s+Restri[\s\S]*?</p>",
    ]
    for pattern in header_patterns:
        template = re.sub(pattern, "", template, count=1, flags=re.IGNORECASE)

    signature_match = re.search(
        r"(<p\b[^>]*>[\s\S]*?S(?:&atilde;|ã|\\u00e3)o\s+Paulo,\s+na\s+data\s+da\s+assinatura\s+digital\.[\s\S]*?</p>)",
        template,
        flags=re.IGNORECASE,
    )
    if signature_match:
        closing = ""
        body_match = re.search(r"</body>\s*</html>\s*$", template, flags=re.IGNORECASE)
        if body_match:
            closing = "\n" + body_match.group(0)
        template = template[: signature_match.end()] + closing

    return template


def ensure_html_utf8(template):
    if re.search(r"<meta[^>]+charset=", template, flags=re.IGNORECASE):
        return re.sub(
            r"<meta[^>]+charset=[\"']?[^\"'>\s]+[\"']?[^>]*>",
            '<meta charset="utf-8">',
            template,
            count=1,
            flags=re.IGNORECASE,
        )
    return re.sub(
        r"<head[^>]*>",
        lambda match: match.group(0) + '\n<meta charset="utf-8">',
        template,
        count=1,
        flags=re.IGNORECASE,
    )


def fill_html(xlsx_path, output_path, sheet_index=0):
    workbook = load_workbook(xlsx_path, data_only=True)
    sheet = workbook["controle"] if "controle" in workbook.sheetnames else workbook.active
    offset = sheet_offset(sheet_index)
    records = load_html_records(sheet, offset)

    if not records:
        raise ValueError(f"Nenhum nome foi encontrado no intervalo D{10 + offset}:D{20 + offset}.")

    result = clean_control_html_template(read_html_template())
    result = replace_html_table(result, records)
    result = replace_html_text(
        result,
        {
            '3º CFO "A"': text(sheet["E4"].value),
            "3º CFO “A”": text(sheet["E4"].value),
            "01 a 30 de abril de 2026": text(sheet["G4"].value),
            "01 a 31 de maio de 2026": text(sheet["G4"].value),
            "HIGOR MACHADO MARQUES": text(cell_value(sheet, "G", 5, offset)),
            "ANDERSON ALVES SILVA": text(cell_value(sheet, "G", 6, offset)),
            "GEAZI DOS SANTOS RODRIGUES": text(cell_value(sheet, "G", 7, offset)),
            'Cad PM - Resp. Restr./Conval./LTS do 3º CFO "A"': text(sheet["F5"].value),
            "Cad PM - Resp. Restr./Conval./LTS do 3º CFO “A”": text(sheet["F5"].value),
            "1º Ten PM - Oficial Responsável": text(sheet["F6"].value),
            "1º Ten PM - Cmt da 1ª Cia": text(sheet["F7"].value),
            "1º Tem PM - Cmt da 1ª Cia Es": text(sheet["F7"].value),
        },
    )
    if offset:
        result = replace_html_text(
            result,
            {
                text(sheet["E4"].value): text(cell_value(sheet, "E", 4, offset)),
                text(sheet["G4"].value): text(cell_value(sheet, "G", 4, offset)),
                text(sheet["F5"].value): text(cell_value(sheet, "F", 5, offset)),
                text(sheet["F6"].value): text(cell_value(sheet, "F", 6, offset)),
                text(sheet["F7"].value): text(cell_value(sheet, "F", 7, offset)),
            },
        )
    result = ensure_html_utf8(result)
    Path(output_path).write_text("\ufeff" + result, encoding="utf-8")
    return len(records)


def choose_file_with_dialog():
    try:
        import tkinter as tk
        from tkinter import filedialog, messagebox
    except Exception as exc:
        raise RuntimeError("Tkinter não está disponível para abrir a janela de seleção.") from exc

    root = tk.Tk()
    root.withdraw()
    xlsx_path = filedialog.askopenfilename(
        title="Selecione a planilha XLSX",
        filetypes=[("Planilhas Excel", "*.xlsx"), ("Todos os arquivos", "*.*")],
    )
    if not xlsx_path:
        return None, None

    default_output = Path(xlsx_path).with_suffix("").with_name(DEFAULT_OUTPUT_NAME)
    output_path = filedialog.asksaveasfilename(
        title="Salvar DOCX preenchido como",
        initialfile=default_output.name,
        defaultextension=".docx",
        filetypes=[("Documentos Word", "*.docx")],
    )
    if not output_path:
        return None, None

    return Path(xlsx_path), Path(output_path)


def parse_args():
    parser = argparse.ArgumentParser(
        description="Gera o controle de Restrição/Convalescença/LTS em DOCX a partir de uma planilha XLSX."
    )
    parser.add_argument("xlsx", nargs="?", help="Caminho da planilha .xlsx")
    parser.add_argument(
        "-o",
        "--output",
        help=f"Caminho do DOCX de saída. Padrão: {DEFAULT_OUTPUT_NAME} na pasta da planilha.",
    )
    return parser.parse_args()


def main():
    args = parse_args()
    used_dialog = not args.xlsx

    try:
        if args.xlsx:
            xlsx_path = Path(args.xlsx).resolve()
            output_path = (
                Path(args.output).resolve()
                if args.output
                else xlsx_path.with_name(DEFAULT_OUTPUT_NAME)
            )
        else:
            xlsx_path, output_path = choose_file_with_dialog()
            if not xlsx_path or not output_path:
                print("Operação cancelada.")
                return

        count, changed = fill_document(xlsx_path, output_path)
        message = (
            f"DOCX gerado: {output_path}\n"
            f"Registros preenchidos: {count}\n"
            f"Campos de cabeçalho/assinatura atualizados: {changed}"
        )
        print(message)

        if used_dialog:
            import tkinter as tk
            from tkinter import messagebox

            root = tk.Tk()
            root.withdraw()
            messagebox.showinfo("Controle gerado", message)
            root.destroy()
    except Exception as exc:
        if used_dialog:
            import tkinter as tk
            from tkinter import messagebox

            root = tk.Tk()
            root.withdraw()
            messagebox.showerror("Erro ao gerar DOCX", str(exc))
            root.destroy()
        raise


if __name__ == "__main__":
    main()
